#!/usr/bin/env python3
"""
Acronym and Abbreviation Extractor

This script extracts acronyms and alphanumeric abbreviations from documents.
Supports .docx, .pdf, and .txt files.

Usage:
    python acronym_extractor.py input_file.docx [output_file.txt]

Requirements:
    pip install python-docx PyPDF2 pandas
"""

import re
import argparse
import sys
from pathlib import Path
from collections import Counter
import pandas as pd

def extract_from_docx(file_path):
    """Extract text from a Word document."""
    try:
        from docx import Document
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"

        return text
    except ImportError:
        print("Error: python-docx not installed. Run: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"Error reading Word document: {e}")
        sys.exit(1)

def extract_from_pdf(file_path):
    """Extract text from a PDF document."""
    try:
        import PyPDF2
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except ImportError:
        print("Error: PyPDF2 not installed. Run: pip install PyPDF2")
        sys.exit(1)
    except Exception as e:
        print(f"Error reading PDF document: {e}")
        sys.exit(1)

def extract_from_txt(file_path):
    """Extract text from a text file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            print(f"Error reading text file: {e}")
            sys.exit(1)
    except Exception as e:
        print(f"Error reading text file: {e}")
        sys.exit(1)

def extract_acronyms_and_abbreviations(text):
    """
    Extract acronyms and alphanumeric abbreviations from text.

    Returns:
        dict: Dictionary with categories of found terms
    """

    # Pattern explanations:
    # 1. Traditional acronyms: 2+ uppercase letters
    # 2. Alphanumeric abbreviations: Mix of letters and numbers
    # 3. Scientific/technical terms: Common patterns in research

    patterns = {
        'traditional_acronyms': r'\b[A-Z]{2,}\b',
        'alphanumeric_abbrev': r'\b[A-Z]*[0-9]+[A-Z]*[0-9]*[A-Z]*\b',
        'mixed_case_abbrev': r'\b[A-Za-z]*[0-9]+[A-Za-z]*\b',
        'scientific_notation': r'\b[A-Z]+[0-9]+[A-Z]*\b'
    }

    results = {}

    for category, pattern in patterns.items():
        matches = re.findall(pattern, text)
        # Filter out common false positives
        filtered_matches = []

        for match in matches:
            # Skip if too short (less than 2 chars)
            if len(match) < 2:
                continue

            # Skip if it's just numbers
            if match.isdigit():
                continue

            # Skip common words that might match patterns
            common_words = {'THE', 'AND', 'FOR', 'ARE', 'BUT', 'NOT', 'YOU', 'ALL', 'CAN', 'HER', 'WAS', 'ONE', 'OUR', 'HAD', 'BUT', 'HAS', 'HIS', 'HIM', 'HER', 'SHE', 'HEP', 'GET', 'USE', 'MAN', 'NEW', 'NOW', 'WAY', 'MAY', 'SAY'}
            if match.upper() in common_words:
                continue

            # Skip single letters followed by numbers that might be page references
            if re.match(r'^[A-Z][0-9]+$', match) and len(match) <= 4:
                # This might be a page reference like "A1", "B23", etc.
                # Include it but flag it
                pass

            filtered_matches.append(match)

        results[category] = filtered_matches

    return results

def combine_results(results):
    """Combine all results without deduplication."""
    all_terms = []

    for category, terms in results.items():
        all_terms.extend(terms)

    return all_terms

def save_results(all_terms, output_file, detailed=False):
    """Save results to file in various formats."""

    if detailed:
        # Create detailed output with categories
        output_path = Path(output_file)
        base_name = output_path.stem

        # Save as CSV for detailed analysis
        csv_file = output_path.parent / f"{base_name}_detailed.csv"

        # Create DataFrame
        data = []
        for i, term in enumerate(all_terms):
            data.append({
                'index': i + 1,
                'term': term,
                'length': len(term),
                'has_numbers': bool(re.search(r'\d', term)),
                'all_caps': term.isupper()
            })

        df = pd.DataFrame(data)
        df.to_csv(csv_file, index=False)
        print(f"Detailed results saved to: {csv_file}")

    # Save simple list
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("EXTRACTED ACRONYMS AND ABBREVIATIONS\n")
        f.write("=" * 50 + "\n\n")
        f.write(f"Total terms found: {len(all_terms)}\n\n")

        f.write("ALL TERMS (IN ORDER FOUND):\n")
        f.write("-" * 30 + "\n")

        for i, term in enumerate(all_terms, 1):
            f.write(f"{i:5d}. {term}\n")

def main():
    parser = argparse.ArgumentParser(
        description="Extract acronyms and abbreviations from documents",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    python acronym_extractor.py document.docx
    python acronym_extractor.py document.pdf output.txt
    python acronym_extractor.py document.txt results.txt --detailed
        """
    )

    parser.add_argument('input_file', help='Input document file (.docx, .pdf, or .txt)')
    parser.add_argument('output_file', nargs='?', help='Output file (optional, defaults to input_acronyms.txt)')
    parser.add_argument('--detailed', action='store_true', help='Generate detailed CSV output with analysis')

    args = parser.parse_args()

    # Validate input file
    input_path = Path(args.input_file)
    if not input_path.exists():
        print(f"Error: Input file '{args.input_file}' not found.")
        sys.exit(1)

    # Determine output file
    if args.output_file:
        output_file = args.output_file
    else:
        output_file = input_path.stem + "_acronyms.txt"

    print(f"Processing: {args.input_file}")
    print("Extracting text...")

    # Extract text based on file type
    file_extension = input_path.suffix.lower()

    if file_extension == '.docx':
        text = extract_from_docx(input_path)
    elif file_extension == '.pdf':
        text = extract_from_pdf(input_path)
    elif file_extension in ['.txt', '.text']:
        text = extract_from_txt(input_path)
    else:
        print(f"Error: Unsupported file type '{file_extension}'. Supported types: .docx, .pdf, .txt")
        sys.exit(1)

    print("Extracting acronyms and abbreviations...")

    # Extract terms
    results = extract_acronyms_and_abbreviations(text)

    # Combine results
    all_terms = combine_results(results)

    print(f"Found {len(all_terms)} total terms")

    # Save results
    save_results(all_terms, output_file, args.detailed)

    print(f"Results saved to: {output_file}")

    if args.detailed:
        try:
            print("Detailed CSV analysis also generated.")
        except ImportError:
            print("Note: Install pandas for detailed CSV output: pip install pandas")

    # Show first 10 results as preview
    print("\nFirst 10 terms found:")
    print("-" * 30)
    for i, term in enumerate(all_terms[:10], 1):
        print(f"{i:2d}. {term}")

if __name__ == "__main__":
    main()
